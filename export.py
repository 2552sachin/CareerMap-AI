"""
export.py - DOCX and PDF export helpers for AI Career Suite
Author: Sachin Chaudhary

Generates Word (.docx) and PDF (.pdf) files from plain-text AI output
and structured analysis results.
"""

import io
import re
from datetime import datetime


def _clean_markdown(text):
    """Strip markdown fences and normalize for document output."""
    if not text:
        return ""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:[a-zA-Z]*)\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _split_lines(text):
    return [ln.rstrip() for ln in text.splitlines() if ln.strip() != ""]


def text_to_docx(title, body, subtitle=None):
    """Return .docx bytes from a title + plain-text body."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x2A, 0x3A)

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x8A, 0x94, 0xA6)
        run.italic = True

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).runs[0].font.size = Pt(8)

    doc.add_paragraph("")

    cleaned = _clean_markdown(body)
    for line in _split_lines(cleaned):
        is_header = line.startswith("#")
        is_bullet = line.startswith(("-", "*", "•"))
        stripped = line.lstrip("#").lstrip("-*• ").rstrip()
        if not stripped:
            continue
        if is_header:
            level = min(line.count("#"), 3)
            try:
                doc.add_heading(stripped, level=max(1, level))
            except Exception:
                doc.add_paragraph(stripped)
        elif is_bullet:
            doc.add_paragraph(stripped, style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def analysis_to_docx(result):
    """Return .docx bytes for a structured resume-analysis result dict."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = f"Resume Analysis — {result.get('job_title','N/A')}"
    if result.get("company_name"):
        title += f" @ {result['company_name']}"
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x2A, 0x3A)

    sub = doc.add_paragraph()
    run = sub.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Provider: {result.get('provider','N/A')}  |  Model: {result.get('model','N/A')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x8A, 0x94, 0xA6)

    doc.add_paragraph("")

    scores = doc.add_paragraph()
    scores.add_run("Scores:  ").bold = True
    scores.add_run(
        f"ATS {result.get('ats_score',0)}/100  |  "
        f"Match {result.get('match_score',0)}/100  |  "
        f"Interview {result.get('hire_probability',0)}/100"
    )

    doc.add_heading("AI Verdict", level=1)
    doc.add_paragraph(result.get("overall_summary", "N/A"))

    def add_list(heading, items, style_name="List Bullet"):
        if not items:
            return
        doc.add_heading(heading, level=2)
        for item in items:
            doc.add_paragraph(str(item), style=style_name)

    add_list("Matched Skills", result.get("matched_skills", []))
    add_list("Missing Skills", result.get("missing_skills", []))
    add_list("ATS Keywords to Add", result.get("keyword_suggestions", []))
    add_list("Strengths", result.get("strengths", []))
    add_list("Improvements", result.get("improvements", []))
    add_list("Quick Wins", result.get("quick_wins", []))
    add_list("Recruiter Red Flags", result.get("red_flags", []))

    if result.get("salary_insight"):
        doc.add_heading("Salary Insight", level=2)
        doc.add_paragraph(result["salary_insight"])
    if result.get("experience_gap"):
        doc.add_heading("Experience Gap", level=2)
        doc.add_paragraph(result["experience_gap"])
    if result.get("education_match"):
        doc.add_heading("Education Match", level=2)
        doc.add_paragraph(result["education_match"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def text_to_pdf(title, body, subtitle=None):
    """Return .pdf bytes from a title + plain-text body using fpdf2."""
    from fpdf import FPDF

    cleaned = _clean_markdown(body)
    pdf = FPDF(format="A4")
    pdf.set_margins(left=15, top=18, right=15)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    w = pdf.epw

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(30, 42, 58)
    pdf.multi_cell(w, 10, title)

    if subtitle:
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(138, 148, 166)
        pdf.multi_cell(w, 5, subtitle)
        pdf.ln(1)

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(138, 148, 166)
    pdf.multi_cell(w, 5, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    pdf.ln(3)
    pdf.set_text_color(30, 42, 58)

    for line in cleaned.splitlines():
        if not line.strip():
            pdf.ln(3)
            continue
        if line.startswith("#"):
            level = min(line.count("#"), 3)
            size = max(13 - level * 2, 10)
            pdf.set_font("Helvetica", "B", size)
            pdf.set_text_color(59, 93, 191)
            pdf.multi_cell(w, size / 2.6, line.lstrip("# ").rstrip())
            pdf.set_text_color(30, 42, 58)
            pdf.set_font("Helvetica", "", 10)
        elif line.startswith(("-", "*", "\u2022")):
            pdf.set_font("Helvetica", "", 10)
            bullet_text = line.lstrip("-*\u2022 ").rstrip()
            pdf.multi_cell(w, 5.5, f"  - {bullet_text}")
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(w, 5.5, line.rstrip())

    out = pdf.output()
    if isinstance(out, str):
        return out.encode("latin-1", errors="replace")
    if isinstance(out, bytearray):
        return bytes(out)
    return bytes(out)


def analysis_to_pdf(result):
    """Return .pdf bytes for a structured resume-analysis result dict."""
    title = f"Resume Analysis - {result.get('job_title','N/A')}"
    if result.get("company_name"):
        title += f" @ {result['company_name']}"
    subtitle = (
        f"Provider: {result.get('provider','N/A')} | Model: {result.get('model','N/A')}"
    )

    lines = []
    lines.append(
        f"Scores: ATS {result.get('ats_score',0)}/100 | "
        f"Match {result.get('match_score',0)}/100 | "
        f"Interview {result.get('hire_probability',0)}/100"
    )
    lines.append("")
    lines.append("# AI Verdict")
    lines.append(result.get("overall_summary", "N/A"))
    lines.append("")

    def section(heading, items):
        out = []
        if items:
            out.append(f"# {heading}")
            for it in items:
                out.append(f"- {it}")
            out.append("")
        return out

    lines += section("Matched Skills", result.get("matched_skills", []))
    lines += section("Missing Skills", result.get("missing_skills", []))
    lines += section("ATS Keywords to Add", result.get("keyword_suggestions", []))
    lines += section("Strengths", result.get("strengths", []))
    lines += section("Improvements", result.get("improvements", []))
    lines += section("Quick Wins", result.get("quick_wins", []))
    lines += section("Recruiter Red Flags", result.get("red_flags", []))

    if result.get("salary_insight"):
        lines += ["# Salary Insight", result["salary_insight"], ""]
    if result.get("experience_gap"):
        lines += ["# Experience Gap", result["experience_gap"], ""]
    if result.get("education_match"):
        lines += ["# Education Match", result["education_match"], ""]

    return text_to_pdf(title, "\n".join(lines), subtitle=subtitle)